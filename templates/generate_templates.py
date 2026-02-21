"""Generate Newport Wholesalers government response templates.

Creates:
  1. Capability Statement (2-page Word doc)
  2. Sources Sought Response Template (Word doc)
  3. Legitimacy Package Checklist (Word doc)

Usage:
    python templates/generate_templates.py
"""

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

OUTPUT_DIR = Path(__file__).resolve().parent
NAVY = RGBColor(0x1B, 0x2A, 0x4A)
GOLD = RGBColor(0xD4, 0xA8, 0x43)


def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading = cell._element.get_or_add_tcPr()
    shading_elem = shading.makeelement(
        qn("w:shd"),
        {qn("w:fill"): color_hex, qn("w:val"): "clear"},
    )
    shading.append(shading_elem)


def add_styled_table(doc, data, col_widths=None, header_bg="1B2A4A"):
    """Add a styled table with navy header."""
    table = doc.add_table(rows=len(data), cols=len(data[0]))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for r, row_data in enumerate(data):
        for c, val in enumerate(row_data):
            cell = table.cell(r, c)
            cell.text = str(val)
            for p in cell.paragraphs:
                p.style = doc.styles["Normal"]
                for run in p.runs:
                    run.font.size = Pt(10)
                    run.font.name = "Calibri"
                    if r == 0:
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                        run.font.bold = True
            if r == 0:
                set_cell_shading(cell, header_bg)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = w

    return table


# ========== DOCUMENT 1: Capability Statement ==========

def generate_capability_statement():
    """Create a 2-page government capability statement."""
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    # === PAGE 1 ===

    # Company Header
    h = doc.add_heading("NEWPORT WHOLESALERS", level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in h.runs:
        run.font.color.rgb = NAVY
        run.font.size = Pt(28)

    tagline = doc.add_paragraph("Florida's Trusted Food Distribution Partner Since [YEAR]")
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in tagline.runs:
        run.font.color.rgb = GOLD
        run.font.size = Pt(14)
        run.font.italic = True

    # Divider
    doc.add_paragraph("_" * 80).alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Company Overview
    h2 = doc.add_heading("COMPANY OVERVIEW", level=2)
    for run in h2.runs:
        run.font.color.rgb = NAVY

    doc.add_paragraph(
        "Newport Wholesalers is a Florida-based broad-line food distribution company with over "
        "25 years of continuous operations. We serve institutional, commercial, and retail food "
        "buyers throughout Florida and the Southeastern United States. Our operations include "
        "temperature-controlled warehousing, a dedicated delivery fleet, and relationships with "
        "hundreds of food manufacturers and suppliers."
    )

    # Core Competencies
    h2 = doc.add_heading("CORE COMPETENCIES", level=2)
    for run in h2.runs:
        run.font.color.rgb = NAVY

    competencies = [
        "Broad-Line Grocery Distribution — Dry, refrigerated, and frozen product categories",
        "Cold Chain Logistics — Temperature-controlled warehousing and delivery fleet",
        "Institutional Food Supply — Schools, hospitals, correctional facilities, military installations",
        "Emergency Response Capability — Florida-based staging and rapid deployment",
        "Product Sourcing & Procurement — Established relationships with 500+ manufacturers",
    ]
    for comp in competencies:
        p = doc.add_paragraph(comp, style="List Bullet")
        for run in p.runs:
            run.font.size = Pt(10)

    # Key Differentiators
    h2 = doc.add_heading("KEY DIFFERENTIATORS", level=2)
    for run in h2.runs:
        run.font.color.rgb = NAVY

    diffs = [
        "25+ Years Continuous Florida Operations — Provable, auditable track record",
        "American-Owned — American employees, American taxes, American supply chain",
        "Real Infrastructure — Warehouses, trucks, cold chain (verifiable upon request)",
        "Small Business Qualification — NAICS 424410 size standard ($200M)",
        "Florida Home Base — Local presence for FL state, county, and school district contracts",
        "Broad Product Range — Single-vendor convenience for food service operations",
    ]
    for diff in diffs:
        p = doc.add_paragraph(diff, style="List Bullet")
        for run in p.runs:
            run.font.size = Pt(10)

    # NAICS Codes
    h2 = doc.add_heading("NAICS CODES", level=2)
    for run in h2.runs:
        run.font.color.rgb = NAVY

    naics_data = [
        ["NAICS Code", "Description"],
        ["424410", "General Line Grocery Merchant Wholesalers (PRIMARY)"],
        ["424420", "Packaged Frozen Food Merchant Wholesalers"],
        ["424430", "Dairy Product Merchant Wholesalers"],
        ["424440", "Poultry and Poultry Product Merchant Wholesalers"],
        ["424450", "Confectionery Merchant Wholesalers"],
        ["424460", "Fish and Seafood Merchant Wholesalers"],
        ["424470", "Meat and Meat Product Merchant Wholesalers"],
        ["424480", "Fresh Fruit and Vegetable Merchant Wholesalers"],
        ["424490", "Other Grocery and Related Products Merchant Wholesalers"],
        ["722310", "Food Service Contractors"],
    ]
    add_styled_table(doc, naics_data, col_widths=[Inches(1.5), Inches(5.5)])

    # Page break
    doc.add_page_break()

    # === PAGE 2 ===

    # Registration & Certifications
    h2 = doc.add_heading("REGISTRATION & CERTIFICATIONS", level=2)
    for run in h2.runs:
        run.font.color.rgb = NAVY

    reg_data = [
        ["Item", "Status", "ID / Details"],
        ["SAM.gov Registration", "[PENDING/ACTIVE]", "UEI: [TO BE ASSIGNED]"],
        ["CAGE Code", "[PENDING/ACTIVE]", "[TO BE ASSIGNED]"],
        ["DUNS Number", "[ACTIVE/MIGRATED]", "[NUMBER]"],
        ["FL Business Registration", "Active", "Filed [YEAR], [XX]+ years continuous"],
        ["MyFloridaMarketPlace (MFMP)", "[PENDING/ACTIVE]", "Vendor ID: [NUMBER]"],
        ["SQF / HACCP Certification", "[STATUS]", "[CERTIFYING BODY]"],
        ["FDA Facility Registration", "[STATUS]", "FEI: [NUMBER]"],
        ["Small Business Self-Certification", "Eligible", "Under $200M size standard (424410)"],
    ]
    add_styled_table(doc, reg_data, col_widths=[Inches(2.5), Inches(1.5), Inches(3)])

    # Past Performance
    h2 = doc.add_heading("PAST PERFORMANCE HIGHLIGHTS", level=2)
    for run in h2.runs:
        run.font.color.rgb = NAVY

    doc.add_paragraph(
        "[To be populated as government contracts are won. In the interim, include "
        "institutional/commercial client references that demonstrate similar capabilities:]"
    ).italic = True

    perf_data = [
        ["Client", "Scope", "Value", "Period", "Reference Available"],
        ["[Major Institutional Client]", "Broad-line food supply", "$[XXX]K/yr", "[YEARS]", "Yes"],
        ["[School/Hospital/Facility]", "Weekly food delivery", "$[XX]K/yr", "[YEARS]", "Yes"],
        ["[Regional Chain/Account]", "Full product catalog supply", "$[XXX]K/yr", "[YEARS]", "Yes"],
    ]
    add_styled_table(doc, perf_data)

    # Facility & Fleet
    h2 = doc.add_heading("FACILITY & FLEET", level=2)
    for run in h2.runs:
        run.font.color.rgb = NAVY

    facility_items = [
        "Primary Warehouse: [ADDRESS], Florida — [XX,XXX] sq ft",
        "Temperature Zones: Ambient, Refrigerated (34-38F), Frozen (-10F to 0F)",
        "Delivery Fleet: [XX] trucks (refrigerated and dry)",
        "Service Area: Florida statewide, Southeast US regional",
        "Operating Hours: [HOURS]",
        "Emergency Response: 24-48 hour rapid deployment capability",
    ]
    for item in facility_items:
        p = doc.add_paragraph(item, style="List Bullet")
        for run in p.runs:
            run.font.size = Pt(10)

    doc.add_paragraph("[INSERT FACILITY PHOTO PLACEHOLDER]").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("[INSERT FLEET PHOTO PLACEHOLDER]").alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Contact
    h2 = doc.add_heading("CONTACT INFORMATION", level=2)
    for run in h2.runs:
        run.font.color.rgb = NAVY

    contact_info = [
        "Newport Wholesalers",
        "[Street Address]",
        "[City], Florida [ZIP]",
        "",
        "Government Sales Contact: [NAME]",
        "Phone: [PHONE]",
        "Email: [EMAIL]",
        "Website: [URL]",
    ]
    for line in contact_info:
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.size = Pt(10)

    path = OUTPUT_DIR / "capability_statement_template.docx"
    doc.save(str(path))
    print(f"  Capability Statement: {path}")
    return path


# ========== DOCUMENT 2: Sources Sought Response Template ==========

def generate_sources_sought_template():
    """Create a Sources Sought response template."""
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Header
    h = doc.add_heading("SOURCES SOUGHT RESPONSE", level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in h.runs:
        run.font.color.rgb = NAVY

    # Reference block
    ref_data = [
        ["Field", "Value"],
        ["Notice Number", "[SAM.GOV NOTICE NUMBER]"],
        ["Notice Title", "[TITLE FROM SAM.GOV]"],
        ["Issuing Agency", "[AGENCY NAME]"],
        ["Response Date", "[DATE]"],
        ["Responding Company", "Newport Wholesalers"],
        ["UEI", "[UNIQUE ENTITY ID]"],
        ["CAGE Code", "[CAGE CODE]"],
        ["NAICS Code", "424410 — General Line Grocery Merchant Wholesalers"],
        ["Business Size", "Small Business"],
        ["Point of Contact", "[NAME, TITLE, PHONE, EMAIL]"],
    ]
    add_styled_table(doc, ref_data, col_widths=[Inches(2), Inches(4.5)])

    doc.add_paragraph("")

    # Section 1
    h2 = doc.add_heading("1. COMPANY OVERVIEW", level=2)
    for run in h2.runs:
        run.font.color.rgb = NAVY

    doc.add_paragraph(
        "Newport Wholesalers is a Florida-based broad-line food distribution company with "
        "over 25 years of continuous operations. We are an American-owned small business "
        "specializing in institutional food supply for schools, government facilities, "
        "correctional institutions, and healthcare systems.\n\n"
        "Our capabilities include temperature-controlled warehousing (ambient, refrigerated, "
        "and frozen), a dedicated delivery fleet serving all of Florida, and established "
        "relationships with 500+ food manufacturers and suppliers."
    )

    # Section 2
    h2 = doc.add_heading("2. RELEVANT CAPABILITIES", level=2)
    for run in h2.runs:
        run.font.color.rgb = NAVY

    doc.add_paragraph(
        "[CUSTOMIZE THIS SECTION FOR EACH NOTICE — address the specific capabilities "
        "requested in the Sources Sought notice. Below are template bullets:]"
    ).italic = True

    capabilities = [
        "Product Range: [Describe relevant product categories — dry goods, refrigerated, "
        "frozen, fresh produce, dairy, meat, beverages, etc.]",
        "Delivery Capability: [Describe delivery frequency, geographic coverage, vehicle types, "
        "cold chain maintenance]",
        "Volume Capacity: [Describe ability to handle the anticipated order volume]",
        "Quality Assurance: [Describe food safety certifications — HACCP, SQF, FDA registration]",
        "Emergency Response: [If relevant — describe rapid deployment capability for "
        "disaster/emergency situations]",
    ]
    for cap in capabilities:
        doc.add_paragraph(cap, style="List Bullet")

    # Section 3
    h2 = doc.add_heading("3. PAST PERFORMANCE", level=2)
    for run in h2.runs:
        run.font.color.rgb = NAVY

    doc.add_paragraph(
        "[Include relevant past performance. For early-stage government work, include "
        "institutional/commercial references that demonstrate similar capabilities:]"
    ).italic = True

    perf_data = [
        ["Client", "Scope of Work", "Value", "Period", "Contact"],
        ["[Client 1]", "[Description]", "$[XX]K", "[Dates]", "[Name, Phone]"],
        ["[Client 2]", "[Description]", "$[XX]K", "[Dates]", "[Name, Phone]"],
        ["[Client 3]", "[Description]", "$[XX]K", "[Dates]", "[Name, Phone]"],
    ]
    add_styled_table(doc, perf_data)

    # Section 4
    h2 = doc.add_heading("4. DELIVERY & LOGISTICS", level=2)
    for run in h2.runs:
        run.font.color.rgb = NAVY

    logistics = [
        "Warehouse Location: [ADDRESS], Florida — [XX,XXX] sq ft",
        "Temperature Zones: Ambient, Refrigerated (34-38F), Frozen (-10F to 0F)",
        "Delivery Fleet: [XX] trucks (refrigerated and dry)",
        "Delivery Schedule: [Frequency available — daily, 3x/week, weekly]",
        "Geographic Coverage: Florida statewide; Southeast US by arrangement",
        "Order Lead Time: [XX] hours from order to delivery",
    ]
    for item in logistics:
        doc.add_paragraph(item, style="List Bullet")

    # Section 5
    h2 = doc.add_heading("5. QUESTIONS & CLARIFICATIONS", level=2)
    for run in h2.runs:
        run.font.color.rgb = NAVY

    doc.add_paragraph(
        "[CUSTOMIZE — ask 2-3 thoughtful questions that show you've read the notice "
        "carefully and are thinking about how to serve the agency well. Example questions:]"
    ).italic = True

    questions = [
        "What is the anticipated delivery schedule and frequency?",
        "Are there specific food safety certifications required beyond HACCP?",
        "Is there an opportunity for an Industry Day or facility tour prior to solicitation?",
        "What is the estimated contract duration (base year + option years)?",
    ]
    for q in questions:
        doc.add_paragraph(q, style="List Number")

    # Section 6
    h2 = doc.add_heading("6. INTEREST STATEMENT", level=2)
    for run in h2.runs:
        run.font.color.rgb = NAVY

    doc.add_paragraph(
        "Newport Wholesalers confirms our interest in this requirement and our capability "
        "to perform the work described. We are a qualified small business under NAICS 424410 "
        "and welcome the opportunity to serve [AGENCY NAME].\n\n"
        "We would welcome the opportunity to participate in any Industry Day events, facility "
        "tours, or pre-solicitation conferences associated with this requirement."
    )

    # Section 7
    h2 = doc.add_heading("7. CONTACT INFORMATION", level=2)
    for run in h2.runs:
        run.font.color.rgb = NAVY

    contact = [
        "Newport Wholesalers",
        "[Name], [Title]",
        "Phone: [PHONE]",
        "Email: [EMAIL]",
        "[Address], Florida [ZIP]",
        "UEI: [NUMBER] | CAGE: [CODE]",
    ]
    for line in contact:
        doc.add_paragraph(line)

    # Footer note
    doc.add_paragraph("")
    p = doc.add_paragraph(
        "INSTRUCTIONS FOR USE: Replace all [BRACKETED] text with actual information. "
        "Customize Section 2 (Capabilities) and Section 5 (Questions) for each specific "
        "Sources Sought notice. Still Mind Creative will prepare customized drafts — "
        "Newport reviews and approves before submission."
    )
    p.italic = True
    for run in p.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    path = OUTPUT_DIR / "sources_sought_response_template.docx"
    doc.save(str(path))
    print(f"  Sources Sought Template: {path}")
    return path


# ========== DOCUMENT 3: Legitimacy Package Checklist ==========

def generate_legitimacy_checklist():
    """Create a legitimacy package checklist."""
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Header
    h = doc.add_heading("LEGITIMACY PACKAGE CHECKLIST", level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in h.runs:
        run.font.color.rgb = NAVY

    tagline = doc.add_paragraph(
        "Newport Wholesalers — Government Vendor Verification Documentation"
    )
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in tagline.runs:
        run.font.color.rgb = GOLD
        run.font.italic = True

    doc.add_paragraph(
        "Purpose: In the current post-fraud-crackdown environment, government procurement "
        "officers verify vendor legitimacy before awarding contracts. This package provides "
        "immediate, verifiable proof that Newport Wholesalers is a real, established, "
        "trustworthy business. Attach relevant components to every proposal and capability statement."
    )

    doc.add_paragraph("")

    # Checklist table
    data = [
        ["#", "Document", "Purpose", "Source", "Status"],
        ["1", "Government Capability Statement\n(2 pages)",
         "Standard government 'business card' showing competencies, NAICS codes, registrations",
         "Still Mind Creative creates\n(template provided)", "[ ] Not Started\n[ ] Draft\n[ ] Complete"],
        ["2", "SAM.gov Registration Confirmation",
         "Proves active federal registration; includes UEI and CAGE code",
         "SAM.gov (30-45 day process)", "[ ] Not Started\n[ ] In Progress\n[ ] Active"],
        ["3", "Florida Business Registration",
         "Proves 25+ year history; shows founding date and continuous operation",
         "Newport provides copy from\nFL Division of Corporations", "[ ] Not Started\n[ ] Requested\n[ ] Obtained"],
        ["4", "Facility Photos (3-5)\n- Exterior with address visible\n- Warehouse interior\n- Cold storage\n- Loading dock",
         "Proves real infrastructure;\nprocurement officers verify addresses",
         "Newport takes photos\n(professional quality preferred)", "[ ] Not Started\n[ ] Scheduled\n[ ] Obtained"],
        ["5", "Fleet Photos (2-3)\n- Branded trucks\n- Refrigerated units",
         "Proves delivery capability;\nshows scale and branding",
         "Newport takes photos", "[ ] Not Started\n[ ] Scheduled\n[ ] Obtained"],
        ["6", "Bank Reference Letter",
         "Proves financial stability;\nshows long-term banking relationship",
         "Newport requests from bank\n(1-2 week turnaround)", "[ ] Not Started\n[ ] Requested\n[ ] Obtained"],
        ["7", "Insurance Certificates\n- General Liability\n- Auto\n- Workers' Comp\n- Product Liability",
         "Proves adequate coverage;\nrequired for most government contracts",
         "Newport requests from\ninsurance broker", "[ ] Not Started\n[ ] Requested\n[ ] Obtained"],
        ["8", "Food Safety Certifications\n- HACCP plan\n- SQF Level 2 (if held)\n- FDA registration",
         "Proves food safety compliance;\nmay be required in solicitations",
         "Newport provides copies\n(or initiates certification)", "[ ] Not Started\n[ ] In Progress\n[ ] Obtained"],
        ["9", "Customer Reference Letters (3-5)\nFrom institutional clients",
         "Third-party validation;\nproves reliable service delivery",
         "Newport requests from\ntop institutional clients", "[ ] Not Started\n[ ] Requested\n[ ] Obtained"],
        ["10", "Product Catalog / Line Card",
         "Shows breadth of product offering;\ndemonstrates single-vendor convenience",
         "Newport provides or\nSMC creates from product data", "[ ] Not Started\n[ ] Draft\n[ ] Complete"],
        ["11", "Company Organizational Chart",
         "Shows management structure;\nproves real staff (not a shell company)",
         "Newport provides", "[ ] Not Started\n[ ] Draft\n[ ] Complete"],
        ["12", "Tax Compliance Certificate\n(FL Dept of Revenue)",
         "Proves state tax compliance;\nrequired for FL state contracts",
         "FL Dept of Revenue website", "[ ] Not Started\n[ ] Requested\n[ ] Obtained"],
    ]
    add_styled_table(doc, data, col_widths=[
        Inches(0.4), Inches(2), Inches(1.8), Inches(1.5), Inches(1)])

    doc.add_paragraph("")

    # Priority notes
    h2 = doc.add_heading("PRIORITY ORDER", level=2)
    for run in h2.runs:
        run.font.color.rgb = NAVY

    priority = [
        "WEEK 1: Items 1 (Capability Statement), 2 (SAM.gov registration — start immediately), "
        "3 (FL Business Registration — quick to obtain)",
        "WEEK 2: Items 4-5 (Photos — schedule a photo session), 6 (Bank letter — request now, "
        "takes 1-2 weeks), 7 (Insurance certs — request from broker)",
        "WEEK 3-4: Items 8 (Food safety certs), 9 (Reference letters — request from 5 clients, "
        "expect 3 back), 10-12 (Supporting documents)",
    ]
    for item in priority:
        doc.add_paragraph(item, style="List Bullet")

    # Why this matters
    h2 = doc.add_heading("WHY THIS MATTERS NOW", level=2)
    for run in h2.runs:
        run.font.color.rgb = NAVY

    doc.add_paragraph(
        "In 2025-2026, the federal government suspended 25% of 8(a) certified contractors "
        "(1,091 firms) for fraudulent small business claims. The $6.8 billion in False Claims "
        "Act recoveries set a new record. Procurement officers are now actively verifying that "
        "vendors are real businesses with real operations.\n\n"
        "Newport's 25-year history IS the competitive advantage — but only if it's packaged "
        "in a way that's easy for procurement officers to verify. This legitimacy package "
        "makes verification effortless and positions Newport as exactly the kind of vendor "
        "that agencies are now seeking."
    )

    path = OUTPUT_DIR / "legitimacy_package_checklist.docx"
    doc.save(str(path))
    print(f"  Legitimacy Checklist: {path}")
    return path


def main():
    print("Generating document templates...")
    generate_capability_statement()
    generate_sources_sought_template()
    generate_legitimacy_checklist()
    print(f"\nAll templates saved to: {OUTPUT_DIR}")


if __name__ == "__main__":
    main()
